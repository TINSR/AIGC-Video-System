from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "TWO_WEEK_MVP_IMPLEMENTATION.md"
OUTPUT = ROOT / "TWO_WEEK_MVP_IMPLEMENTATION.docx"


def set_east_asia(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_east_asia(run, "Microsoft YaHei")
    run.font.size = Pt(9.5)
    run.bold = bold
    if bold:
        run.font.color.rgb = RGBColor(31, 77, 120)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2F3")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "电商场景 AIGC 带货视频生成系统 - 两周 MVP 实现文档"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_east_asia(run, "Microsoft YaHei")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 116, 139)


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code.rstrip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8.5)


def add_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [c.strip().replace("`", "") for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return
    data_rows = [rows[0]] + rows[2:]
    col_count = max(len(row) for row in data_rows)
    table = doc.add_table(rows=len(data_rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for r_idx, row in enumerate(data_rows):
        for c_idx in range(col_count):
            text = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(table.cell(r_idx, c_idx), text, bold=(r_idx == 0))
            if r_idx == 0:
                shade_cell(table.cell(r_idx, c_idx), "F2F4F7")
    doc.add_paragraph()


def add_paragraph_with_inline_code(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(31, 77, 120)
        else:
            run = p.add_run(part)
            set_east_asia(run, "Microsoft YaHei")


def build() -> None:
    doc = Document()
    apply_styles(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []
    table_buffer: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buffer.append(line)
            i += 1
            next_line = lines[i] if i < len(lines) else ""
            if not (next_line.strip().startswith("|") and next_line.strip().endswith("|")):
                add_table(doc, table_buffer)
                table_buffer = []
            continue

        stripped = line.strip()
        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(stripped[2:])
            set_east_asia(run, "Microsoft YaHei")
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(11, 37, 69)
            run.bold = True
        elif stripped.startswith("## "):
            doc.add_paragraph(stripped[3:], style="Heading 1")
        elif stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 2")
        elif stripped.startswith("#### "):
            doc.add_paragraph(stripped[5:], style="Heading 3")
        elif stripped.startswith("> "):
            add_paragraph_with_inline_code(doc, stripped[2:])
        elif stripped.startswith("- "):
            add_paragraph_with_inline_code(doc, stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            add_paragraph_with_inline_code(doc, re.sub(r"^\d+\. ", "", stripped), style="List Number")
        else:
            add_paragraph_with_inline_code(doc, stripped)

        i += 1

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
